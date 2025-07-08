import openai
import json
import re
import os
import base64
from typing import List, Dict, Any
from colorama import Fore, Style
from config import Config
from rag_engine import RAGEngine
from document_processor import DocumentProcessor
from pdf_processor import PDFProcessor

def extract_json_from_response(text):
    """
    Extracts the first legal JSON block from the AI return, removing redundant descriptions.
    """
    if not text or not text.strip():
        return ""
    text = re.sub(r'```(?:json)?', '', text, flags=re.IGNORECASE).strip()
    match = re.search(r'\{[\s\S]*\}', text)
    if match:
        return match.group(0)
    return text

class AIClient:
    def __init__(self):
        self.client = openai.OpenAI(
            api_key=Config.OPENAI_API_KEY,
            base_url=Config.OPENAI_BASE_URL
        )
        self.model = Config.OPENAI_MODEL
        self.rag_engine = RAGEngine()
        self.doc_processor = DocumentProcessor()
        self.pdf_processor = PDFProcessor()

    def analyze_empty_fields_by_index(self, document_content: str) -> Dict[str, Any]:
        """Analyze all indexed fields in the document, decide which need to be filled, and restore content for those that do not."""
        prompt = f"""
        You are given a document with all table cells labeled with an [index] (e.g., [1], [2], etc.).
        For each cell:
        - If the cell does NOT need to be filled (e.g., it already contains valid content or should be left as is such as the row header, column header and title), return its original content (remove the [index] tag and restore the cell text).
        - If the cell SHOULD be filled (e.g., it is a placeholder, missing, or needs user input), return its index, a short description of what should be filled, and a suggested content type (e.g., text, date, number).
        
        Please return a JSON object with two keys:
        {{
          "fields_to_fill": [
            {{
              "index": 1,
              "description": "The meaning of this field, e.g., Gender",
              "suggested_content_type": "Type of content (e.g., text, date, number)"
            }}
          ],
          "restored_cells": [
            {{
              "index": 2,
              "restored_content": "Original cell content without [2]"
            }}
          ]
        }}
        
        Document content:
        {document_content}

        Only return valid JSON.
        """
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a smart document understanding assistant, good at understanding labeled fields in tables."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1
            )
            print("AI raw response:", response)
            result = response.choices[0].message.content
            
            print(f"AI response content: {result}")
            
            if not result or not result.strip():
                print("Warning: AI returned empty response")
                return {"fields_to_fill": [], "restored_cells": []}
            
            if result.strip().startswith("```"):
                match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", result, re.IGNORECASE)
                if match:
                    result = match.group(1)
            
            print(f"Processed result: {result}")
            
            result = extract_json_from_response(result)
            parsed_result = json.loads(result)
            print(f"Successfully parsed JSON: {parsed_result}")
            return parsed_result
            
        except json.JSONDecodeError as e:
            print(f"JSON decode error: {e}")
            print(f"Failed to parse: {result}")
            return {"fields_to_fill": [], "restored_cells": []}
        except Exception as e:
            print(f"Error analyzing fields by index: {e}")
            return {"fields_to_fill": [], "restored_cells": []}

    def search_with_rag(self, field_info: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        if self.rag_engine.get_index_stats()["document_count"] == 0:
            print("Txt knowledge base is empty, cannot perform RAG search.")
            return []

        all_results = []
        for field in field_info:
            index = field.get("index")
            desc = field.get("description", "")
            print(f"\n{Fore.CYAN} Field [{index}] - {desc}{Style.RESET_ALL}")

            results = self.rag_engine.semantic_search([field], top_k=3)

            if results:
                for i, r in enumerate(results, 1):
                    content = r.get("content", "").strip()
                    score = r.get("similarity_score", 0)
                    print(f"{Fore.GREEN}  {i}. Score: {score:.3f}{Style.RESET_ALL}")
                    print(f"     Content: {content}")
            else:
                print(f"{Fore.RED}  No RAG result found for this field{Style.RESET_ALL}")

            field["rag_evidence"] = results if results else []

            all_results.extend(results if results else [])

        return all_results

    def fill_document_with_rag(self, field_info: List[Dict[str, Any]], user_data: List[Dict[str, Any]]) -> Dict[str, Any]:
        rag_context = self._build_rag_context(user_data)
        prompt = f"""
        Based on the following field information (each with an index) and user data, generate appropriate content for each field.

        Fields to be filled:
        {json.dumps(field_info, ensure_ascii=False, indent=2)}

        User data (RAG retrieval results):
        {rag_context}

        Please return in JSON:
        {{
            "field_answers": [
                {{
                    "index": 1,
                    "content": "content to fill",
                    "confidence": "high/medium/low",
                    "source": "data source"
                }}
            ]
        }}
        """
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a professional document filling assistant using RAG."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2
            )
            print("AI raw response:", response)
            result = response.choices[0].message.content
            if result.strip().startswith("```"):
                match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", result, re.IGNORECASE)
                if match:
                    result = match.group(1)
            print(f"fill_document_with_rag result: {result}")
            result = extract_json_from_response(result)
            return json.loads(result)
        except Exception as e:
            print(f"Error generating fill content with AI: {e}")
            return {"field_answers": []}

    def _build_rag_context(self, user_data: List[Dict[str, Any]]) -> str:
        if not user_data:
            return "No relevant user data."
        return "\n".join(user.get("content", "") for user in user_data if user.get("content"))

    def get_rag_stats(self) -> Dict[str, Any]:
        return self.rag_engine.get_index_stats()

    def update_rag_index(self, new_documents: List[Dict[str, Any]]):
        self.rag_engine.update_index(new_documents)

    def fill_document(self, field_info: List[Dict[str, Any]], user_data: List[Dict[str, Any]]) -> Dict[str, Any]:
        return self.fill_document_with_rag(field_info, user_data)

    def final_fill_decision(self, all_fields: list, described_fields: list) -> dict:
        """
        Let the LLM decide, for each cell, whether to fill it (and with what content) or restore its original content,
        based on the cell's content, description, and RAG evidence.
        """
        index_to_desc = {f["index"]: f for f in described_fields}
        merged_fields = []
        for field in all_fields:
            idx = field["index"]
            desc_info = index_to_desc.get(idx, {})
            merged = field.copy()
            merged["description"] = desc_info.get("description", "")
            merged["suggested_content_type"] = desc_info.get("suggested_content_type", "")
            merged["rag_evidence"] = desc_info.get("rag_evidence", [])
            merged_fields.append(merged)

        prompt = f"""
        You are an intelligent document filling assistant. For each table cell (with index), you are given:
        - The cell's current content (may include an [index] tag)
        - A semantic description of what the cell means
        - RAG evidence (knowledge base search results, may be empty)

        For each cell:
        - If there is enough RAG evidence to fill the cell, output the content to fill. (If there is placeholder in the cell, keep it and append the content to fill after it)
        - If there is no relevant RAG evidence, or the cell should not be filled, restore the cell's original content (remove the [index] tag and keep the original text).

        Return a JSON object with two keys:
        {{
          "filled_cells": [
            {{ "index": 1, "content": "..." }}
          ],
          "restored_cells": [
            {{ "index": 2, "restored_content": "..." }}
          ]
        }}

        Input fields:
        {json.dumps(merged_fields, ensure_ascii=False, indent=2)}

        Only return valid JSON.
        """
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a professional document filling assistant using RAG."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2
            )
            print("AI raw response:", response)
            result = response.choices[0].message.content
            if result.strip().startswith("```"):
                match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", result, re.IGNORECASE)
                if match:
                    result = match.group(1)
            result = extract_json_from_response(result)
            return json.loads(result)
        except Exception as e:
            print(f"Error in final fill decision: {e}")
            return {"filled_cells": [], "restored_cells": []}

    def analyze_empty_fields_with_images(self, document_content: str, page_images: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Analyze empty fields in document using both images and text content"""
        
        messages = [
            {
                "role": "system", 
                "content": "You are a professional document understanding assistant, skilled at analyzing fields in tables. Please carefully examine the document images and text content to identify which fields need to be filled."
            }
        ]
        
        for page_info in page_images:
            page_num = page_info['page_number']
            image_path = page_info['image_path']
            
            with open(image_path, 'rb') as img_file:
                image_base64 = base64.b64encode(img_file.read()).decode('utf-8')
            
            messages.append({
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"This is page {page_num} of the document:"
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{image_base64}"
                        }
                    }
                ]
            })
        
        prompt = f"""
        Based on the above document images and the following text content, analyze all table cells with indices (e.g., [1], [2], etc.).

        For each cell:
        - If the cell does NOT need to be filled (e.g., it already contains valid content or should be left as is such as row headers, column headers and title), return its original content (remove the [index] tag and restore the cell text).
        - If the cell SHOULD be filled (e.g., it is a placeholder, missing, or needs user input), return its index, a short description of what should be filled, and a suggested content type (e.g., text, date, number).
        - Also include standalone text lines if they are followed by a colon (:) but have no content — these also count as fields that need to be filled.
        - For multi-row sections, please assume that every row's cells should be filled by default, don't just take the first line, but note that headers must be restored.
        - Note that no cell can belong to both two categories.

        Please return a JSON object with two keys:
        {{
          "fields_to_fill": [
            {{
              "index": 1,
              "description": "The meaning of this field, e.g., Gender",
              "suggested_content_type": "Type of content (e.g., text, date, number)"
            }}
          ],
          "restored_cells": [
            {{
              "index": 2,
              "restored_content": "Original cell content without [2]"
            }}
          ]
        }}

        Document text content:
        {document_content}

        Please carefully analyze the table structure and content in the images, combined with text information to make accurate judgments.
        Only return valid JSON.
        """
        
        messages.append({
            "role": "user",
            "content": prompt
        })
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=0.1,
                max_tokens=20000
            )
            
            print("AI raw response:", response)
            result = response.choices[0].message.content
            
            print(f"AI response content: {result}")
            
            if not result or not result.strip():
                print("Warning: AI returned empty response")
                return {"fields_to_fill": [], "restored_cells": []}
            
            if result.strip().startswith("```"):
                match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", result, re.IGNORECASE)
                if match:
                    result = match.group(1)
            
            print(f"Processed result: {result}")
            
            result = extract_json_from_response(result)
            parsed_result = json.loads(result)
            print(f"Successfully parsed JSON: {parsed_result}")
            return parsed_result
            
        except json.JSONDecodeError as e:
            print(f"JSON decode error: {e}")
            print(f"Failed to parse: {result}")
            return {"fields_to_fill": [], "restored_cells": []}
        except Exception as e:
            print(f"Error analyzing fields with images: {e}")
            return {"fields_to_fill": [], "restored_cells": []}

    def split_text_with_llm(self, text: str, max_chunk_length: int = 300) -> list:
        """
        Use LLM to split long text into semantically complete and contextually coherent chunks of appropriate length.
        Returns: list of chunk strings
        """
        prompt = f"""
        Please intelligently split the following text into several semantically complete and contextually coherent chunks. Each chunk should preferably be a complete sentence or paragraph, and the length should be appropriate (suggested: no more than {max_chunk_length} characters per chunk).
        Output format: Only return a JSON array, each element is a string representing a chunk.
        50~100 words per chunk is recommended.
        Don't modify the original text.
        Original text:
        {text}
        """
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an intelligent assistant skilled at splitting long text into semantically complete chunks."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1
            )
            result = response.choices[0].message.content
            print("LLM chunking raw response:", result)
            if result.strip().startswith("```"):
                match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", result, re.IGNORECASE)
                if match:
                    result = match.group(1)
            result = extract_json_from_response(result)
            chunks = json.loads(result)
            if isinstance(chunks, list):
                return [str(c).strip() for c in chunks if str(c).strip()]
            else:
                return []
        except Exception as e:
            print(f"Error in split_text_with_llm: {e}")
            return []

    def extract_knowledge_from_file(self, file_path: str) -> str:
        """
        Extract knowledge information from files, supporting multiple file formats
        """
        if not os.path.exists(file_path):
            print(f"{Fore.RED}File not found: {file_path}{Style.RESET_ALL}")
            return ""
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext in ['.txt', '.md']:
                return self._extract_from_text_file(file_path)
            elif file_ext in ['.doc', '.docx']:
                return self._extract_from_word_file(file_path)
            elif file_ext == '.pdf':
                return self._extract_from_pdf_file(file_path)
            else:
                print(f"{Fore.YELLOW}Unsupported file format: {file_ext}{Style.RESET_ALL}")
                return ""
        except Exception as e:
            print(f"{Fore.RED}Error extracting knowledge from file: {e}{Style.RESET_ALL}")
            return ""

    def _extract_from_text_file(self, file_path: str) -> str:
        """Extract knowledge from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return self._process_text_content(content, "text file")
        except Exception as e:
            print(f"{Fore.RED}Failed to read text file: {e}{Style.RESET_ALL}")
            return ""

    def _extract_from_word_file(self, file_path: str) -> str:
        """Extract knowledge from Word file"""
        try:
            content = self._extract_all_text_from_docx(file_path)
            return self._process_text_content(content, "Word document")
        except Exception as e:
            print(f"{Fore.RED}Failed to read Word file: {e}{Style.RESET_ALL}")
            return ""

    def _extract_all_text_from_docx(self, file_path: str) -> str:
        """
        Extract all text content from Word document including paragraphs, tables, headers, footers
        Args:
            file_path: Path to the Word document
        Returns:
            All text content as string
        """
        try:
            from docx import Document
            doc = Document(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content.append(cell_text)
                    if row_content:
                        table_content.append(' | '.join(row_content))
                if table_content:
                    content_parts.append('\n'.join(table_content))
            
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(f"[Header] {paragraph.text.strip()}")
            
            for section in doc.sections:
                footer = section.footer
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(f"[Footer] {paragraph.text.strip()}")
            
            full_content = '\n\n'.join(content_parts)
            return full_content
            
        except Exception as e:
            print(f"Error extracting text from Word document: {e}")
            return ""

    def _extract_from_pdf_file(self, file_path: str) -> str:
        """Extract knowledge from PDF file"""
        try:
            content = self.pdf_processor.extract_text_from_pdf(file_path)
            return self._process_text_content(content, "PDF document")
        except Exception as e:
            print(f"{Fore.RED}Failed to read PDF file: {e}{Style.RESET_ALL}")
            return ""

    def _process_text_content(self, content: str, file_type: str) -> str:
        """Use AI to process text content and extract structured knowledge"""
        if not content.strip():
            return ""
        
        prompt = f"""
        You are a professional document knowledge extraction assistant. Please extract useful knowledge information from the following {file_type} content:

        Requirements:
        1. Extract all important information points, including but not limited to: personal information, company information, project information, dates, amounts, addresses, etc.
        2. Maintain the completeness and accuracy of information
        3. Organize related information into structured text
        4. Remove irrelevant formatting information and duplicate content
        5. Preserve key numbers, dates, names and other specific information

        Document content:
        {content}

        Please return the extracted knowledge text directly, without additional formatting instructions.
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a professional document knowledge extraction assistant, skilled at extracting structured information from various documents."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=4000
            )
            
            result = response.choices[0].message.content
            print(f"{Fore.GREEN}✓ Successfully extracted knowledge from {file_type}{Style.RESET_ALL}")
            return result.strip()
            
        except Exception as e:
            print(f"{Fore.RED}AI processing of text content failed: {e}{Style.RESET_ALL}")
            return content.strip()

    def build_rag_from_files(self, file_paths: List[str]) -> bool:
        """
        Build RAG knowledge base from multiple files
        """
        print(f"{Fore.CYAN}Starting to build RAG knowledge base from files...{Style.RESET_ALL}")
        
        all_knowledge = []
        success_count = 0
        
        for i, file_path in enumerate(file_paths, 1):
            print(f"{Fore.YELLOW}[{i}/{len(file_paths)}] Processing file: {os.path.basename(file_path)}{Style.RESET_ALL}")
            
            knowledge = self.extract_knowledge_from_file(file_path)
            
            if knowledge:
                chunks = self.split_text_with_llm(knowledge, max_chunk_length=300)
                
                if chunks:
                    documents = [{'id': f"{os.path.basename(file_path)}_{j}", 'content': chunk} 
                               for j, chunk in enumerate(chunks)]
                    all_knowledge.extend(documents)
                    success_count += 1
                    print(f"{Fore.GREEN}✓ Successfully processed: {os.path.basename(file_path)} ({len(chunks)} knowledge chunks){Style.RESET_ALL}")
                else:
                    print(f"{Fore.YELLOW}⚠ File processing completed but no valid knowledge chunks generated: {os.path.basename(file_path)}{Style.RESET_ALL}")
            else:
                print(f"{Fore.RED}✗ File processing failed: {os.path.basename(file_path)}{Style.RESET_ALL}")
        
        if all_knowledge:
            self.update_rag_index(all_knowledge)
            print(f"{Fore.GREEN}✓ RAG knowledge base construction completed! Processed {success_count}/{len(file_paths)} files, generated {len(all_knowledge)} knowledge chunks{Style.RESET_ALL}")
            
            stats = self.get_rag_stats()
            print(f"{Fore.CYAN}RAG statistics: {stats}{Style.RESET_ALL}")
            return True
        else:
            print(f"{Fore.RED}✗ Failed to extract valid knowledge from any files{Style.RESET_ALL}")
            return False

    def extract_knowledge_from_file_legacy(self, file_type: str, file_base64: str) -> str:
        """
        Legacy version of knowledge extraction from file (for compatibility)
        """
        prompt = f"""
        You are a professional document understanding assistant, skilled at extracting knowledge from files.
        This is a {file_type} file with base64 encoding {file_base64}, please extract the knowledge information from the file.
        
        Requirements:
        1. Extract all important information points
        2. Maintain the completeness and accuracy of information
        3. Organize related information into structured text
        4. Remove irrelevant formatting information
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are a professional document understanding assistant, skilled at extracting knowledge from files."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1
            )
            result = response.choices[0].message.content
            print("AI raw response:", result)
            if result.strip().startswith("```"):
                match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", result, re.IGNORECASE)
                if match:
                    result = match.group(1)
            return result
        except Exception as e:
            print(f"Error in extract_knowledge_from_file_legacy: {e}")
            return ""