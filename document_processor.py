import os
import re
import time
from typing import List, Dict, Any, Tuple
from docx import Document
from docx.oxml.shared import OxmlElement, qn
import openpyxl
from openpyxl.styles import PatternFill
from openpyxl.cell import MergedCell
from config import Config
import subprocess
import shutil

class DocumentProcessor:
    def __init__(self):
        self.highlight_color = Config.HIGHLIGHT_COLOR

    def extract_document_content(self, file_path: str) -> str:
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext == '.docx':
            return self._extract_docx_content(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self._extract_excel_content(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    def _extract_docx_content(self, file_path: str) -> str:
        doc = Document(file_path)
        content = []
        for table in doc.tables:
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    row_content.append(cell.text.strip())
                content.append(' | '.join(row_content))
        return '\n'.join(content)

    def _extract_excel_content(self, file_path: str) -> str:
        wb = openpyxl.load_workbook(file_path)
        content = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            content.append(f"=== Sheet: {sheet_name} ===")
            
            for row in ws.iter_rows():
                row_content = []
                for cell in row:
                    row_content.append(str(cell.value).strip())
                content.append(' | '.join(row_content))
        wb.close()
        return '\n'.join(content)

    def _save_cell_format(self, cell):
        """Saving cell formatting information"""
        format_info = {
            'paragraphs': []
        }
        
        for paragraph in cell.paragraphs:
            para_format = {
                'alignment': paragraph.alignment,
                'runs': []
            }
            
            for run in paragraph.runs:
                run_format = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': run.font.color.rgb if run.font.color.rgb else None
                }
                para_format['runs'].append(run_format)
            
            format_info['paragraphs'].append(para_format)
        
        return format_info

    def _apply_cell_format(self, cell, format_info, new_text):
        """Apply saved formatting to cells"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        
        if not cell.paragraphs:
            cell.add_paragraph()
        
        if format_info['paragraphs']:
            para_format = format_info['paragraphs'][0]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = para_format['alignment']
            
            if para_format['runs']:
                run_format = para_format['runs'][0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.text = new_text
                run.bold = run_format['bold']
                run.italic = run_format['italic']
                run.underline = run_format['underline']
                if run_format['font_name']:
                    run.font.name = run_format['font_name']
                if run_format['font_size']:
                    run.font.size = run_format['font_size']
                if run_format['font_color']:
                    run.font.color.rgb = run_format['font_color']
            else:
                paragraph.text = new_text
        else:
            cell.paragraphs[0].text = new_text

    def find_and_number_all_fields_docx(self, file_path: str) -> Tuple[List[Dict[str, Any]], str]:
        """为Word文档所有字段添加索引"""
        doc = Document(file_path)
        all_fields = []
        field_index = 1
        key_to_index = {}

        for table_index, table in enumerate(doc.tables):
            for row_index, row in enumerate(table.rows):
                for col_index, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if re.search(r'\[\d+\]$', cell_text):
                        continue
                    
                    original_format = self._save_cell_format(cell)
                    
                    key = id(cell._tc)
                    key_to_index[key] = field_index
                    
                    if cell_text:
                        new_text = f"{cell_text} [{field_index}]"
                    else:
                        new_text = f"[{field_index}]"
                    
                    self._apply_cell_format(cell, original_format, new_text)
                    
                    all_fields.append({
                        'index': field_index,
                        'type': 'table_cell',
                        'table_index': table_index,
                        'row_index': row_index,
                        'col_index': col_index,
                        'text': cell_text,
                        'original_format': original_format,
                        'context': f'Table {table_index+1}, Row {row_index+1}, Col {col_index+1}'
                    })
                    field_index += 1

        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_path = os.path.join(Config.MID_DIR, f"{base_name}_numbered_{int(time.time())}.docx")
        doc.save(output_path)
        return all_fields, output_path

    def find_and_number_all_fields_excel(self, file_path: str) -> Tuple[List[Dict[str, Any]], str]:
        """为Excel文档所有字段添加索引"""
        wb = openpyxl.load_workbook(file_path)
        all_fields = []
        field_index = 1

        for sheet_index, sheet_name in enumerate(wb.sheetnames):
            ws = wb[sheet_name]
            
            for row_index, row in enumerate(ws.iter_rows(), 1):
                for col_index, cell in enumerate(row, 1):
                    # 跳过合并单元格中的非主单元格
                    if isinstance(cell, MergedCell):
                        print(f"Skipping merged cell at Row {row_index}, Col {col_index}")
                        continue
                    
                    cell_value = str(cell.value) if cell.value is not None else ""
                    
                    # 检查是否已经有索引标记
                    if re.search(r'\[\d+\]$', cell_value):
                        continue
                    
                    # 添加索引标记
                    try:
                        if cell_value.strip():
                            cell.value = f"{cell_value} [{field_index}]"
                        else:
                            cell.value = f"[{field_index}]"
                        
                        all_fields.append({
                            'index': field_index,
                            'type': 'excel_cell',
                            'sheet_index': sheet_index,
                            'sheet_name': sheet_name,
                            'row_index': row_index,
                            'col_index': col_index,
                            'text': cell_value.strip(),
                            'context': f'Sheet {sheet_name}, Row {row_index}, Col {col_index}'
                        })
                        field_index += 1
                        
                    except Exception as e:
                        print(f"Warning: Cannot modify cell at Row {row_index}, Col {col_index}: {e}")
                        continue

        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_path = os.path.join(Config.MID_DIR, f"{base_name}_numbered_{int(time.time())}.xlsx")
        wb.save(output_path)
        wb.close()
        return all_fields, output_path

    def find_and_number_all_fields(self, file_path: str) -> Tuple[List[Dict[str, Any]], str]:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return self.find_and_number_all_fields_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self.find_and_number_all_fields_excel(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    def fill_document(self, file_path: str, field_answers: List[Dict[str, Any]]) -> str:
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext == '.docx':
            return self._fill_docx_document(file_path, field_answers)
        elif file_ext in ['.xlsx', '.xls']:
            return self._fill_excel_document(file_path, field_answers)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    def _fill_docx_document(self, file_path: str, field_answers: List[Dict[str, Any]]) -> str:
        """填充Word文档"""
        doc = Document(file_path)

        for ans in field_answers:
            try:
                table_idx = ans["table_index"]
                row_idx = ans["row_index"]
                col_idx = ans["col_index"]
                content = ans["content"]
                cell = doc.tables[table_idx].cell(row_idx, col_idx)
                
                original_format = ans.get("original_format")
                if original_format:
                    self._apply_cell_format(cell, original_format, str(content))
                else:
                    cell.text = str(content)
            except Exception as e:
                print(f"Failed to fill cell index {ans.get('index')}: {e}")

        original_filename = os.path.basename(file_path)
        name_without_ext = original_filename.split("_numbered")[0].split("_highlighted")[0].split("_restored")[0].replace(".docx", "")
        filled_filename = f"{name_without_ext}.docx"
        filled_path = os.path.join(Config.OUTPUT_DIR, filled_filename)
        doc.save(filled_path)
        return filled_path

    def _fill_excel_document(self, file_path: str, field_answers: List[Dict[str, Any]]) -> str:
        """填充Excel文档"""
        wb = openpyxl.load_workbook(file_path)

        for ans in field_answers:
            try:
                sheet_name = ans["sheet_name"]
                row_idx = ans["row_index"]
                col_idx = ans["col_index"]
                content = ans["content"]
                
                ws = wb[sheet_name]
                cell = ws.cell(row=row_idx, column=col_idx)
                
                # 检查是否是合并单元格
                if isinstance(cell, MergedCell):
                    print(f"Warning: Cannot fill merged cell at Row {row_idx}, Col {col_idx}")
                    continue
                    
                cell.value = str(content)
                
            except Exception as e:
                print(f"Failed to fill cell index {ans.get('index')}: {e}")

        original_filename = os.path.basename(file_path)
        name_without_ext = original_filename.split("_numbered")[0].split("_highlighted")[0].split("_restored")[0]
        name_without_ext = os.path.splitext(name_without_ext)[0]
        filled_filename = f"{name_without_ext}.xlsx"
        filled_path = os.path.join(Config.OUTPUT_DIR, filled_filename)
        wb.save(filled_path)
        wb.close()
        return filled_path

    def restore_cells_content_from_indexed_docx(self, file_path: str, restored_cells: List[Dict[str, Any]]) -> str:
        """恢复Word文档单元格内容"""
        doc = Document(file_path)
        for cell_info in restored_cells:
            idx = cell_info.get("index")
            content = cell_info.get("restored_content", "")
            original_format = cell_info.get("original_format")
            found = False
            for table_index, table in enumerate(doc.tables):
                for row_index, row in enumerate(table.rows):
                    for col_index, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text.endswith(f"[{idx}]") or cell_text == f"[{idx}]":
                            if original_format:
                                self._apply_cell_format(cell, original_format, content)
                            else:
                                cell.text = content
                            found = True
                            break
                    if found:
                        break
                if found:
                    break
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_path = os.path.join(Config.MID_DIR, f"{base_name}_restored_{int(time.time())}.docx")
        doc.save(output_path)
        return output_path

    def to_cmd_path(self, path):
        return os.path.normpath(path).replace('\\', '/')

    def convert_doc_to_docx(self, doc_path, docx_path):
        """Convert doc to docx using libreoffice"""
        output_dir = os.path.dirname(docx_path)
        doc_path_cmd = self.to_cmd_path(doc_path)
        docx_path_cmd = self.to_cmd_path(docx_path)
        output_dir_cmd = self.to_cmd_path(output_dir)
        soffice_path = self.to_cmd_path(os.path.join('C:', 'Program Files', 'LibreOffice', 'program', 'soffice.exe'))
        subprocess.run([
            soffice_path,
            '--headless',
            '--convert-to', 'docx',
            doc_path_cmd,
            '--outdir', output_dir_cmd
        ], check=True)
        base = os.path.splitext(os.path.basename(doc_path))[0]
        generated_docx = os.path.join(output_dir, base + '.docx')
        generated_docx = self.to_cmd_path(generated_docx)
        if generated_docx != docx_path_cmd:
            shutil.move(generated_docx, docx_path_cmd)
        return docx_path_cmd
    
    def restore_cells_content_from_indexed_excel(self, file_path: str, restored_cells: List[Dict[str, Any]]) -> str:
        """恢复Excel文档单元格内容"""
        wb = openpyxl.load_workbook(file_path)
        
        for cell_info in restored_cells:
            idx = cell_info.get("index")
            content = cell_info.get("restored_content", "")
            found = False
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows():
                    for cell in row:
                        # 跳过合并单元格
                        if isinstance(cell, MergedCell):
                            continue
                            
                        cell_value = str(cell.value) if cell.value is not None else ""
                        if cell_value.endswith(f"[{idx}]") or cell_value == f"[{idx}]":
                            try:
                                cell.value = content
                                found = True
                                break
                            except Exception as e:
                                print(f"Warning: Cannot restore cell with index {idx}: {e}")
                                continue
                    if found:
                        break
                if found:
                    break
        
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_path = os.path.join(Config.MID_DIR, f"{base_name}_restored_{int(time.time())}.xlsx")
        wb.save(output_path)
        wb.close()
        return output_path

    def restore_cells_content_from_indexed(self, file_path: str, restored_cells: List[Dict[str, Any]]) -> str:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return self.restore_cells_content_from_indexed_docx(file_path, restored_cells)
        elif file_ext in ['.xlsx', '.xls']:
            return self.restore_cells_content_from_indexed_excel(file_path, restored_cells)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

